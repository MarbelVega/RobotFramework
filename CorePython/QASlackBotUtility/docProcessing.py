from docx import Document
import csv
import glob
import zipfile
import os
from xml.dom import minidom
import shutil
import time

def create_blank_docx_template_with_name(name, extension='docx'):
    doc = Document()
    doc.save(f'{name}.{extension}')

def create_standard_docx_template(name, data, extension='docx'):
    doc = Document()
    for row in data:
        doc.add_paragraph(row)
    doc.save(f'{name}.{extension}')

def compare_csv_reports(before, after):
    # compares two report csv files for firm specific reports, and tells which we broke.
    new_csv_data = []
    # read the two files to compare
    with open(before, newline='') as f:
        reader = csv.reader(f)
        data_before = list(reader)
    with open(after, newline='') as f:
        reader = csv.reader(f)
        data_after = list(reader)
    # go through each row on the newest report, and compare to the previous
    for line in data_after:
        has_been_broken = False
        # Take care of headers
        if line[0].lower() == 'firm id':
            print("Header Row")
            new_csv_data.append([line[0], line[1], line[2], line[3], line[4], line[5], line[6], line[7], line[8], "Has Been Broken"])
        # Otherwise, since it's not a header row, go through the report
        else:
            for row in data_before:
                # while the second column is all we care about, we are matching on the first and second columns to determine
                #   that we are looking at the correct report results
                if line[0] == row[0] and line[1] == row[1] and line[3].lower() == 'fail' and row[3].lower() == 'pass':
                    has_been_broken = True
            # For now, we are logging a Yes if we broke it, and nothing if we did not.
            if has_been_broken:
                new_csv_data.append([line[0], line[1], line[2], line[3], line[4], line[5], line[6], line[7], line[8], 'Yes'])
            else:
                new_csv_data.append([line[0], line[1], line[2], line[3], line[4], line[5], line[6], line[7], line[8], ' '])
    with open("results.csv", "w", newline='') as f:
        writer = csv.writer(f)
        writer.writerows(new_csv_data)


# Unzips an IDML file to the unzip directory for further processing
def unzip_idml_template(idml_template, unzip_directory):
    with zipfile.ZipFile(idml_template, 'r') as z:
        z.extractall(unzip_directory)

# zips back up an IDML file to be an IDML file
def zip_idml_template(idml_template_name, idml_template_directory):
    zipf = zipfile.ZipFile(idml_template_name, 'w', zipfile.ZIP_DEFLATED)
    exclude_dir = len(idml_template_directory)  # with idml templates, we need to exclude the outermost directory
    for root, _, files in os.walk(idml_template_directory):
        folder = root[exclude_dir:]  # exclude the base path
        for file in files:
            zipf.write(os.path.join(root, file), os.path.join(folder, file))
    zipf.close()

# goes through the idml stories and returns the content text for further processing
def get_content_from_idml_template(path, idml_template):
    expected_idml_directory = f"{path}{idml_template.replace('idml', '')}"
    unzip_idml_template(f"{path}{idml_template}", expected_idml_directory)
    word_list = []
    for _, _, files in os.walk(f"{expected_idml_directory}/Stories"):
        for file in files:
            doc = minidom.parse(f"{expected_idml_directory}/Stories/{file}")
            contents = doc.getElementsByTagName("Content")
            for content in contents:
                word_list.append(content.firstChild.nodeValue)
    shutil.rmtree(expected_idml_directory)
    return word_list

def get_text_from_downloaded_docx(path, docname):
    doc = Document(f'{path}{docname.replace(" ", "")}.docx')
    rows = []
    for row in doc.paragraphs:
        rows.append(row.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                rows.append(cell.text)
    
    return rows

def does_exported_excel_have_an_error(path):
    # Assumes that there will only be one csv
    try:
        # Wait 30 seconds max for a file to download... this is usually enough time.
        loop_break = 0
        while loop_break < 30:
            files = glob.glob(f"{path}\\*.csv")
            print(files)
            if len(files) == 0:
                print(f"File not found, sleeping #{loop_break} for 1 second...")
                time.sleep(1)
                loop_break += 1
            else:
                break
        truthy = False
        for file in files:
            with open(file) as csv_file:
                csv_reader = csv.reader(csv_file, delimiter=',')
                for row in csv_reader:
                    if 'The error has been logged and tech support has been automatically informed.' in row:
                        truthy=True
    except:
        print("There was an unexplainable error.")
        truthy = 'Unknown'
    finally:
        return truthy

def cleanup_directory(path):
    files = glob.glob(f"{path}\\*.csv")        
    for file in files:
        os.remove(file)


def verify_no_variables_in_list(word_list):
    truthy = True

    for row in word_list:
        if "%" in row:
            truthy = False
    
    return truthy

def verify_template_variables_replaced(word_list, doc_type):
    truthy = False
    csvf = None

    if doc_type.lower() == 'resume':
        csvf = 'C:\\Users\\rlewis\\Documents\\robotframework\\Cosential\\Publisher Templates\\simpleresumedata.csv'
    if doc_type.lower() == 'project profile':
        csvf = 'C:\\Users\\rlewis\\Documents\\robotframework\\Cosential\\Publisher Templates\\simpleprojectprofiledata.csv'
    if doc_type.lower() == 'project list':
        csvf = 'C:\\Users\\rlewis\\Documents\\robotframework\\Cosential\\Publisher Templates\\simpleprojectlistdata.csv'
    if doc_type.lower() == 'opportunity profile':
        csvf = 'C:\\Users\\rlewis\\Documents\\robotframework\\Cosential\\Publisher Templates\\simpleopportunityprofiledata.csv'

    with open(csvf, newline='') as f:
        reader = csv.reader(f)
        data = list(reader)

    for row in data:
        for r in word_list:
            if row[0].lower() in r.lower():
                truthy = True
        if truthy == False:
            return False
        truthy = False    
    return True
